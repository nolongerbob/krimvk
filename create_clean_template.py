#!/usr/bin/env python3
"""Create a clean Word template with properly formatted placeholders."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new document
doc = Document()

# Add title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("ЗАЯВЛЕНИЕ\nо выдаче технических условий")
run.bold = True
run.font.size = Pt(14)

# Add main content
doc.add_paragraph()

# Personal info section
p = doc.add_paragraph()
p.add_run("От: {{FIO}}, дата рождения {{BirthDate}}, паспорт {{Passport}}, ")
p.add_run("ИНН {{inn}}, СНИЛС {{snils}}")

doc.add_paragraph()

# Address section
p = doc.add_paragraph()
p.add_run("Адрес регистрации: {{registration address}}")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Почтовый адрес: {{mail address}}")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Телефон: {{phone}}, Email: {{email}}")

doc.add_paragraph()
doc.add_paragraph()

# Request section
p = doc.add_paragraph()
p.add_run("Прошу выдать технические условия для присоединения к централизованным системам водоснабжения и (или) водоотведения.")

doc.add_paragraph()

# Basis
p = doc.add_paragraph()
p.add_run("Основание обращения с запросом о выдаче технических условий:")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("{{owner or trusted person}}")

doc.add_paragraph()

# Reason
p = doc.add_paragraph()
p.add_run("В связи с: {{reason}}")

doc.add_paragraph()

# Object info
p = doc.add_paragraph()
p.add_run("Тип объекта: {{object_type}}")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Адрес объекта: {{place address}}")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Тип подключения: {{type of connection}}")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Необходимые виды ресурсов: {{resourse_type}}")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Информация об объекте: {{object information}}")

doc.add_paragraph()

# Date
p = doc.add_paragraph()
p.add_run("Планируемый срок ввода в эксплуатацию: {{date_start}}")

doc.add_paragraph()

# Result delivery
p = doc.add_paragraph()
p.add_run("Результат предоставления услуги прошу направить: {{where to receive result}}")

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Signature
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run("{{fio}}")
p.add_run("\n_______________(подпись)")

# Save the document
output_path = "public/documents/tu-template-clean.docx"
doc.save(output_path)
print(f"✅ Clean template created: {output_path}")
